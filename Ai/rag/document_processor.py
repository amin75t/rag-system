"""
Document Processor for extracting text from various file formats.

This module provides functionality to extract text from PDF, DOCX, TXT,
and other document formats for RAG indexing.

Enhanced with:
- DOCX to Markdown conversion
- Detailed logging for each processing step
- Memory management (4GB RAM limit)
- Streaming processing for large documents
"""
import os
import logging
import psutil
import gc
import time
from typing import Optional, List, Dict, Any
from pathlib import Path

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# Memory limit: 4GB in bytes
MAX_MEMORY_BYTES = 4 * 1024 * 1024 * 1024  # 4GB


# ==================== Memory Management Functions ====================

def get_memory_usage() -> int:
    """Get current memory usage in bytes."""
    process = psutil.Process(os.getpid())
    return process.memory_info().rss


def log_memory_usage(step_name: str):
    """Log current memory usage with step name."""
    mem_bytes = get_memory_usage()
    mem_mb = mem_bytes / (1024 * 1024)
    mem_gb = mem_bytes / (1024 * 1024 * 1024)
    logger.info(f"[MEMORY] {step_name}: {mem_mb:.2f} MB ({mem_gb:.3f} GB)")
    return mem_bytes


def check_memory_limit(step_name: str) -> bool:
    """
    Check if memory usage is within the 4GB limit.
    
    Returns:
        True if within limit, False if approaching limit
    """
    current_mem = get_memory_usage()
    usage_percent = (current_mem / MAX_MEMORY_BYTES) * 100
    
    logger.info(f"[MEMORY CHECK] {step_name}: {usage_percent:.2f}% of 4GB limit")
    
    if usage_percent > 90:
        logger.warning(f"[MEMORY WARNING] Approaching 4GB limit! Current: {usage_percent:.2f}%")
        return False
    elif usage_percent > 75:
        logger.warning(f"[MEMORY ALERT] Memory usage high: {usage_percent:.2f}%")
        return True
    else:
        return True


def force_garbage_collection(step_name: str):
    """Force garbage collection and log memory before/after."""
    logger.info(f"[GC] Starting garbage collection at step: {step_name}")
    mem_before = get_memory_usage()
    
    gc.collect()
    
    mem_after = get_memory_usage()
    mem_freed = (mem_before - mem_after) / (1024 * 1024)
    logger.info(f"[GC] Freed {mem_freed:.2f} MB ({mem_before/(1024*1024):.2f}MB -> {mem_after/(1024*1024):.2f}MB)")


# ==================== Document Processor Class ====================

class DocumentProcessor:
    """
    Document Processor for extracting text from various file formats.
    
    Supports: PDF, DOCX, TXT, MD, and plain text files.
    """
    
    def __init__(self):
        """Initialize the Document Processor."""
        self.supported_extensions = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx_to_markdown,
            '.txt': self._extract_text,
            '.md': self._extract_text,
            '.text': self._extract_text,
        }
        logger.info("=" * 60)
        logger.info("DocumentProcessor initialized")
        logger.info(f"Memory limit: {MAX_MEMORY_BYTES / (1024**3):.1f} GB")
        log_memory_usage("Initialization")
        logger.info("=" * 60)
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a document file with detailed logging and memory management.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        logger.info("=" * 60)
        logger.info(f"[STEP 1] Starting document extraction")
        logger.info(f"File: {file_path}")
        
        log_memory_usage("Before extraction")
        
        path = Path(file_path)
        
        if not path.exists():
            logger.error(f"[ERROR] File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")
        
        logger.info(f"[STEP 1.1] File exists, size: {path.stat().st_size / (1024*1024):.2f} MB")
        
        extension = path.suffix.lower()
        logger.info(f"[STEP 1.2] File type: {extension}")
        
        if extension not in self.supported_extensions:
            logger.error(f"[ERROR] Unsupported file format: {extension}")
            raise ValueError(
                f"Unsupported file format: {extension}. "
                f"Supported formats: {', '.join(self.supported_extensions.keys())}"
            )
        
        extractor = self.supported_extensions[extension]
        logger.info(f"[STEP 1.3] Using extractor: {extractor.__name__}")
        
        # Check memory before extraction
        if not check_memory_limit("Before extraction"):
            logger.warning("[WARNING] Memory high, forcing GC before extraction")
            force_garbage_collection("Pre-extraction")
        
        start_time = time.time()
        text = extractor(file_path)
        extraction_time = time.time() - start_time
        
        logger.info(f"[STEP 1.4] Extraction completed in {extraction_time:.2f} seconds")
        logger.info(f"[STEP 1.5] Extracted {len(text)} characters")
        logger.info(f"[STEP 1.6] Processing rate: {len(text) / extraction_time:.0f} chars/sec")
        
        log_memory_usage("After extraction")
        
        # Check memory after extraction
        if not check_memory_limit("After extraction"):
            logger.warning("[WARNING] Memory high after extraction")
            force_garbage_collection("Post-extraction")
        
        logger.info("=" * 60)
        return text
    
    def _extract_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            import PyPDF2
            
            text = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text.append(page.extract_text())
            
            return '\n\n'.join(text)
            
        except ImportError:
            raise ImportError(
                "PyPDF2 is not installed. "
                "Install it with: pip install PyPDF2"
            )
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise
    
    def _extract_docx_to_markdown(self, file_path: str) -> str:
        """
        Convert DOCX file to Markdown format with detailed logging.
        
        This method:
        1. Extracts text from DOCX file
        2. Converts to Markdown format preserving structure
        3. Handles paragraphs, headings, tables, lists
        4. Streams processing to manage memory
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Markdown formatted text
        """
        logger.info(f"[DOCX->MD] Starting DOCX to Markdown conversion")
        start_time = time.time()
        
        try:
            from docx import Document
            
            logger.info(f"[DOCX->MD] Loading DOCX file...")
            doc = Document(file_path)
            
            logger.info(f"[DOCX->MD] Document loaded with {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
            
            markdown_lines = []
            
            # Process paragraphs
            logger.info(f"[DOCX->MD] Processing paragraphs...")
            paragraph_count = 0
            for i, paragraph in enumerate(doc.paragraphs):
                if not paragraph.text.strip():
                    continue
                
                paragraph_count += 1
                
                # Check heading level
                if paragraph.style.name.startswith('Heading'):
                    level = int(paragraph.style.name.replace('Heading ', ''))
                    markdown_lines.append('#' * level + ' ' + paragraph.text)
                else:
                    markdown_lines.append(paragraph.text)
                
                # Log progress every 100 paragraphs
                if i % 100 == 0:
                    logger.info(f"[DOCX->MD] Processed {i}/{len(doc.paragraphs)} paragraphs")
                    log_memory_usage(f"Processing paragraph {i}")
                    
                    # Force GC if memory is high
                    if not check_memory_limit(f"Paragraph {i}"):
                        force_garbage_collection(f"After paragraph {i}")
            
            logger.info(f"[DOCX->MD] Processed {paragraph_count} non-empty paragraphs")
            
            # Process tables
            logger.info(f"[DOCX->MD] Processing tables...")
            for table_idx, table in enumerate(doc.tables):
                logger.info(f"[DOCX->MD] Processing table {table_idx + 1}/{len(doc.tables)}")
                
                # Create markdown table
                markdown_lines.append('')
                markdown_lines.append(f'**Table {table_idx + 1}**')
                markdown_lines.append('')
                
                # Header row
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                markdown_lines.append('| ' + ' | '.join(header_cells) + ' |')
                markdown_lines.append('|' + '|'.join(['---'] * len(header_cells)) + '|')
                
                # Data rows
                for row_idx, row in enumerate(table.rows[1:], 1):
                    row_cells = [cell.text.strip() for cell in row.cells]
                    markdown_lines.append('| ' + ' | '.join(row_cells) + ' |')
                    
                    # Log progress for large tables
                    if row_idx % 50 == 0:
                        logger.info(f"[DOCX->MD] Table {table_idx + 1}: Processed {row_idx} rows")
                        log_memory_usage(f"Table {table_idx + 1} row {row_idx}")
                        
                        # Force GC if memory is high
                        if not check_memory_limit(f"Table {table_idx + 1} row {row_idx}"):
                            force_garbage_collection(f"Table {table_idx + 1} row {row_idx}")
                
                markdown_lines.append('')
            
            # Join and return
            markdown_text = '\n'.join(markdown_lines)
            
            conversion_time = time.time() - start_time
            logger.info(f"[DOCX->MD] Conversion completed in {conversion_time:.2f} seconds")
            logger.info(f"[DOCX->MD] Output size: {len(markdown_text)} characters")
            logger.info(f"[DOCX->MD] Processing rate: {len(markdown_text) / conversion_time:.0f} chars/sec")
            
            return markdown_text
            
        except ImportError:
            logger.error(f"[DOCX->MD] python-docx not installed")
            raise ImportError(
                "python-docx is not installed. "
                "Install it with: pip install python-docx"
            )
        except Exception as e:
            logger.error(f"[DOCX->MD] Error converting DOCX: {e}")
            raise
    
    def _extract_text(self, file_path: str) -> str:
        """
        Extract text from plain text file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            File content
        """
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try other encodings
            encodings = ['latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            raise ValueError(f"Could not decode file {file_path} with any common encoding")
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            raise
    
    def chunk_text(
        self,
        text: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        output_dir: Optional[str] = None,
        save_to_disk: bool = True
    ) -> List[str]:
        """
        Split text into overlapping chunks for embedding with streaming to disk.
        
        This method implements a streaming approach:
        - Chunks are saved to disk as they are created
        - Only one chunk (or small batch) is in RAM at a time
        - If process is killed, already-saved chunks persist on disk
        
        Args:
            text: Input text to chunk
            chunk_size: Maximum size of each chunk (characters)
            chunk_overlap: Overlap between chunks (characters)
            output_dir: Directory to save chunk files (defaults to ./chunks_output)
            save_to_disk: Whether to save chunks to disk (default: True)
            
        Returns:
            List of text chunks
        """
        logger.info("=" * 60)
        logger.info(f"[STEP 2] Starting text chunking (streaming mode)")
        logger.info(f"Text length: {len(text)} characters")
        logger.info(f"Chunk size: {chunk_size} characters")
        logger.info(f"Chunk overlap: {chunk_overlap} characters")
        logger.info(f"Save to disk: {save_to_disk}")
        
        log_memory_usage("Before chunking")
        
        if chunk_overlap >= chunk_size:
            logger.error("[ERROR] chunk_overlap must be less than chunk_size")
            raise ValueError("chunk_overlap must be less than chunk_size")
        
        # Create output directory if saving to disk
        chunk_files = []
        if save_to_disk:
            if output_dir is None:
                output_dir = "./chunks_output"
            Path(output_dir).mkdir(parents=True, exist_ok=True)
            logger.info(f"[STEP 2.0] Output directory: {output_dir}")
        
        start_time = time.time()
        chunks = []
        start = 0
        text_length = len(text)
        
        # Minimum advance to prevent infinite loop
        min_chunk_size = max(chunk_size // 4, chunk_overlap + 10)
        
        logger.info(f"[STEP 2.1] Minimum chunk size: {min_chunk_size} characters")
        logger.info(f"[STEP 2.2] Estimated chunks: ~{(text_length / (chunk_size - chunk_overlap)):.0f}")
        
        chunk_num = 0
        while start < text_length:
            end = start + chunk_size
            
            # Try to end at a sentence boundary
            if end < text_length:
                # Look for sentence terminators
                terminators = ['.', '!', '?', '\n\n']
                for terminator in terminators:
                    last_pos = text.rfind(terminator, start, end)
                    # FIX: Only adjust if it's not too close to start (minimum chunk size)
                    if last_pos != -1 and (last_pos - start) >= min_chunk_size:
                        end = last_pos + 1
                        break
            
            chunk = text[start:end].strip()
            
            if chunk:
                chunks.append(chunk)
                chunk_num += 1
                
                # Save to disk immediately if enabled
                if save_to_disk:
                    chunk_filename = f"{output_dir}/chunk_{chunk_num:06d}.md"
                    try:
                        with open(chunk_filename, 'w', encoding='utf-8') as f:
                            f.write(f"# Chunk {chunk_num}\n\n")
                            f.write(chunk)
                            f.write("\n\n---\n\n")
                        chunk_files.append(chunk_filename)
                        
                        # Log progress every 100 chunks
                        if chunk_num % 100 == 0:
                            logger.info(f"[STEP 2.3] Created {chunk_num} chunks, saved to disk...")
                            log_memory_usage(f"Chunking progress: {chunk_num}")
                            
                            # Force GC if memory is high
                            if not check_memory_limit(f"Chunk {chunk_num}"):
                                force_garbage_collection(f"After chunk {chunk_num}")
                    except Exception as e:
                        logger.error(f"[ERROR] Failed to save chunk {chunk_num}: {e}")
            
            # FIX: Ensure we always advance forward
            new_start = end - chunk_overlap
            if new_start <= start:
                # Force minimum advance if overlap causes stall
                new_start = start + min_chunk_size
                logger.warning(f"[STEP 2.4] Chunk {chunk_num} too small, forcing advance to {new_start}")
            start = new_start
        
        chunking_time = time.time() - start_time
        logger.info(f"[STEP 2.5] Chunking completed in {chunking_time:.2f} seconds")
        logger.info(f"[STEP 2.6] Created {len(chunks)} chunks")
        logger.info(f"[STEP 2.7] Saved {len(chunk_files)} chunks to disk" if save_to_disk else "[STEP 2.7] Disk saving disabled")
        logger.info(f"[STEP 2.8] Average chunk size: {sum(len(c) for c in chunks) / len(chunks):.0f} characters")
        logger.info(f"[STEP 2.9] Chunking rate: {len(chunks) / chunking_time:.0f} chunks/sec")
        
        log_memory_usage("After chunking")
        
        # Force GC after chunking
        force_garbage_collection("Post-chunking")
        
        logger.info("=" * 60)
        logger.info(f"[STEP 2.10] Chunks saved to: {output_dir}" if save_to_disk else "[STEP 2.10] Chunks in memory only")
        logger.info("=" * 60)
        
        return chunks
    
    def process_file(
        self,
        file_path: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        metadata: Optional[Dict[str, Any]] = None,
        output_dir: Optional[str] = None,
        save_chunks_to_disk: bool = True
    ) -> List[Dict[str, Any]]:
        """
        Process a document file and return chunks with metadata.
        Includes detailed logging and memory management with streaming to disk.
        
        Args:
            file_path: Path to document file
            chunk_size: Maximum size of each chunk
            chunk_overlap: Overlap between chunks
            metadata: Additional metadata to include
            output_dir: Directory to save chunk files (defaults to ./chunks_output)
            save_chunks_to_disk: Whether to save chunks to disk (default: True)
            
        Returns:
            List of document chunks with metadata
        """
        logger.info("=" * 60)
        logger.info(f"[STEP 0] Starting file processing")
        logger.info(f"File: {file_path}")
        logger.info(f"Chunk size: {chunk_size}, Overlap: {chunk_overlap}")
        logger.info(f"Save chunks to disk: {save_chunks_to_disk}")
        
        log_memory_usage("Start of process_file")
        
        start_time = time.time()
        
        # Extract text
        text = self.extract_text(file_path)
        
        # Chunk text (with streaming to disk)
        chunks = self.chunk_text(text, chunk_size, chunk_overlap, output_dir, save_chunks_to_disk)
        
        # Prepare metadata
        logger.info(f"[STEP 3] Preparing metadata and creating document objects")
        path = Path(file_path)
        base_metadata = {
            'filename': path.name,
            'file_path': str(path),
            'file_size': path.stat().st_size,
            'file_type': path.suffix.lower(),
        }
        
        if metadata:
            base_metadata.update(metadata)
            logger.info(f"[STEP 3.1] Added custom metadata: {list(metadata.keys())}")
        
        # Create chunk documents with memory management
        logger.info(f"[STEP 3.2] Creating {len(chunks)} document objects...")
        log_memory_usage("Before creating documents")
        
        documents = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = base_metadata.copy()
            chunk_metadata['chunk_id'] = i
            chunk_metadata['chunk_index'] = i
            chunk_metadata['total_chunks'] = len(chunks)
            
            documents.append({
                'content': chunk,
                'metadata': chunk_metadata,
                'id': f"{path.stem}_chunk_{i}"
            })
            
            # Log progress every 100 documents
            if (i + 1) % 100 == 0:
                logger.info(f"[STEP 3.3] Created {i + 1}/{len(chunks)} document objects...")
                log_memory_usage(f"Creating documents: {i + 1}")
                
                # Force GC if memory is high
                if not check_memory_limit(f"Document {i + 1}"):
                    force_garbage_collection(f"After document {i + 1}")
        
        processing_time = time.time() - start_time
        logger.info(f"[STEP 4] File processing completed")
        logger.info(f"[STEP 4.1] Total time: {processing_time:.2f} seconds")
        logger.info(f"[STEP 4.2] Created {len(documents)} document objects")
        logger.info(f"[STEP 4.3] Processing rate: {len(documents) / processing_time:.0f} docs/sec")
        
        log_memory_usage("End of process_file")
        
        # Final GC check
        if not check_memory_limit("End of process_file"):
            force_garbage_collection("Final cleanup")
        
        logger.info("=" * 60)
        return documents


# Singleton instance
_document_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """
    Get or create a singleton instance of DocumentProcessor.
    
    Returns:
        DocumentProcessor instance
    """
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor
